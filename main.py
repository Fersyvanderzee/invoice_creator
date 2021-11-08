import docx
import datetime

doc = docx.Document('template.docx')
invoice_nr = str(datetime.datetime.now().year) + str(datetime.datetime.now().month)
date_tuple = (str(datetime.datetime.now().day), str(datetime.datetime.now().month), str(datetime.datetime.now().year))
date = '-'.join(date_tuple)

all_paras = doc.paragraphs

i = 0

for para in all_paras:
    print(str(i) + '' + para.text)
    print('------')
    i += 1

all_paras[18].text = f'Factuurnummer: {invoice_nr}'
all_paras[19].text = f'Factuurdatum: {date}'



f = docx.Document()

f.add_paragraph(all_paras[18].text)
f.add_paragraph(all_paras[19].text)

f.save('test.docx')




